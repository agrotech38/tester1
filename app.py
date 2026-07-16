# app.py
"""
Merged PSS + COA Generator
---------------------------
Single Streamlit app that:
  1. Collects one shared set of inputs (type, date, PO ID, container info, batch data)
  2. Fills the PSS template (MOD PSS.docx / FAR PSS.docx)
  3. Fills the COA template (PH LIPL MOD COA.docx / PH LIPL FAR COA.docx)
  4. Merges the two filled .docx files into a single .docx (PSS first, then a
     page break, then COA) using docxcompose
  5. (Optional) Converts the merged .docx to a single PDF via headless
     LibreOffice, if you also want a PDF download

Requires the `docxcompose` package (docx merge) and, only if you use the PDF
button, LibreOffice on the host (see packages.txt for Streamlit Community
Cloud deployment).
"""

import streamlit as st
from datetime import datetime
import re
import os
import io
import subprocess
import tempfile
from docx import Document
from docxcompose.composer import Composer

try:
    from zoneinfo import ZoneInfo
    KOLKATA = ZoneInfo("Asia/Kolkata")
except Exception:
    import pytz
    KOLKATA = pytz.timezone("Asia/Kolkata")

st.set_page_config(page_title="PSS + COA Generator", layout="wide")
st.title("📄 PSS + COA Generator")

# ==================================================================
# Template locations
# ==================================================================
PSS_TEMPLATES = {"MOD": "MOD PSS.docx", "FAR": "FAR PSS.docx"}
COA_TEMPLATES = {"MOD": "PH LIPL MOD COA.docx", "FAR": "PH LIPL FAR COA.docx"}


def find_template(path):
    """Look for the template next to the app, or in /mnt/data as a fallback."""
    candidates = [path, os.path.join("/mnt/data", path), os.path.join("templates", path)]
    for c in candidates:
        if os.path.exists(c):
            return c
    return None


# ==================================================================
# PSS-style replacement: plain whole-paragraph text substitution
# (PSS templates mix tokens like "{{DD/MM/YYYY}}" and bare "DD/MM/YYYY")
# ==================================================================
def replace_in_paragraph_by_text(paragraph, mapping):
    text = paragraph.text
    new_text = text
    for key, val in mapping.items():
        if key in new_text:
            new_text = new_text.replace(key, val)
    if new_text != text:
        paragraph.text = new_text


def replace_text_in_block(block, mapping):
    for paragraph in getattr(block, "paragraphs", []):
        replace_in_paragraph_by_text(paragraph, mapping)
    for table in getattr(block, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_block(cell, mapping)


def apply_pss_replacements(doc, mapping):
    replace_text_in_block(doc, mapping)
    for section in doc.sections:
        try:
            replace_text_in_block(section.header, mapping)
        except Exception:
            pass
        try:
            replace_text_in_block(section.footer, mapping)
        except Exception:
            pass


# ==================================================================
# COA-style replacement: regex "{{KEY}}" matching, preserves run style
# ==================================================================
PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_\-/]+)\s*\}\}")


def normalize_broken_placeholders_in_doc(doc):
    def fix_runs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if "((" in run.text or "))" in run.text:
                    run.text = run.text.replace("((", "{{").replace("))", "}}")

    fix_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fix_runs(cell.paragraphs)
    try:
        for section in doc.sections:
            fix_runs(section.header.paragraphs)
            fix_runs(section.footer.paragraphs)
    except Exception:
        pass


def replace_placeholders_in_paragraph(paragraph, replacements):
    runs = paragraph.runs
    if not runs:
        return

    full_text = ""
    offsets = []
    for run in runs:
        start = len(full_text)
        full_text += run.text
        end = len(full_text)
        offsets.append((run, start, end))

    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    for match in matches:
        key = match.group(1)
        if key not in replacements:
            continue
        replacement_text = str(replacements[key])
        p_start, p_end = match.start(), match.end()

        overlapping = [(r, s, e) for (r, s, e) in offsets if not (e <= p_start or s >= p_end)]
        if not overlapping:
            continue

        first_run, first_s, first_e = overlapping[0]
        last_run, last_s, last_e = overlapping[-1]

        prefix_len = max(0, p_start - first_s)
        prefix = first_run.text[:prefix_len]
        suffix_start_in_last = p_end - last_s
        suffix = last_run.text[suffix_start_in_last:]

        for r, _, _ in overlapping:
            r.text = ""

        new_text = prefix + replacement_text + suffix
        first_run.text = new_text

        try:
            font = first_run.font
            if font.name:
                first_run.font.name = font.name
            if font.size:
                first_run.font.size = font.size
            first_run.font.bold = font.bold
            first_run.font.italic = font.italic
            first_run.font.underline = font.underline
            if font.color and getattr(font.color, "rgb", None) is not None:
                first_run.font.color.rgb = font.color.rgb
        except Exception:
            pass


def apply_coa_replacements(doc, replacements):
    normalize_broken_placeholders_in_doc(doc)

    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, replacements)

    try:
        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_placeholders_in_paragraph(para, replacements)
            for para in section.footer.paragraphs:
                replace_placeholders_in_paragraph(para, replacements)
    except Exception:
        pass


# ==================================================================
# Helpers
# ==================================================================
def fill_template(path, mapping, style_preserving):
    doc = Document(path)
    if style_preserving:
        apply_coa_replacements(doc, mapping)
    else:
        apply_pss_replacements(doc, mapping)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def docx_bytes_to_pdf_bytes(docx_bytes, base_name):
    """Convert docx bytes to pdf bytes using headless LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, f"{base_name}.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                "--convert-to", "pdf", "--outdir", tmpdir, docx_path,
            ],
            capture_output=True,
            timeout=180,
        )

        pdf_path = os.path.join(tmpdir, f"{base_name}.pdf")
        if not os.path.exists(pdf_path):
            stderr = result.stderr.decode(errors="ignore") if result.stderr else "unknown error"
            raise RuntimeError(f"LibreOffice conversion failed for {base_name}: {stderr}")

        with open(pdf_path, "rb") as f:
            return f.read()


def merge_docx(first_docx_bytes, second_docx_bytes):
    """
    Merge two .docx files into one: first_docx_bytes, a page break,
    then the full content of second_docx_bytes (paragraphs, tables,
    styles, images all carried over via docxcompose).
    """
    master = Document(io.BytesIO(first_docx_bytes))
    master.add_page_break()

    composer = Composer(master)
    sub_doc = Document(io.BytesIO(second_docx_bytes))
    composer.append(sub_doc)

    out = io.BytesIO()
    composer.save(out)
    out.seek(0)
    return out.read()


def extract_3digit_code(po_id: str) -> str:
    if not po_id:
        return "000"
    candidate = str(po_id).strip()
    digits = re.findall(r"\d", candidate)
    if len(digits) >= 3:
        return "".join(digits[-3:])
    alnum = re.findall(r"[A-Za-z0-9]", candidate)
    if len(alnum) >= 3:
        return "".join(alnum[-3:]).upper()
    cleaned = re.sub(r"[^A-Za-z0-9]", "", candidate)
    return cleaned[-3:].upper() if len(cleaned) >= 3 else cleaned.zfill(3)


# ==================================================================
# UI
# ==================================================================
now_kolkata = datetime.now(KOLKATA)

doc_type = st.selectbox("Choose Type", ["MOD", "FAR"])

tab_general, tab1, tab2, tab3, tab4 = st.tabs(["General", "Batch 1", "Batch 2", "Batch 3", "Batch 4"])

batches = {"1": {}, "2": {}, "3": {}, "4": {}}

with tab_general:
    date_picker = st.date_input("Date", value=now_kolkata.date())
    po_id = st.text_input("P.O. ID", value="LIPL2026270", key="po_id")
    total_containers = st.number_input("Total Container", min_value=1, step=1, value=1, key="total_containers")
    current_container = st.number_input("Current Container", min_value=1, step=1, value=1, key="current_container")

batch_tabs = {"1": tab1, "2": tab2, "3": tab3, "4": tab4}
for i, tab in batch_tabs.items():
    with tab:
        st.subheader(f"Batch {i}")
        # Shared batch number -> used as PSS {{B<i>}} AND COA BATCH_<i>
        batches[i]["BATCH"] = st.text_input("Batch Number", key=f"batch{i}_label")
        batches[i]["M"] = st.text_input("Moisture", key=f"m{i}")
        if doc_type == "MOD":
            batches[i]["B1V1"] = st.text_input("30min Viscosity", key=f"b{i}v1_mod")
            batches[i]["B1V2"] = st.text_input("60min Viscosity", key=f"b{i}v2_mod")
            batches[i]["PH"] = st.text_input("pH", key=f"ph{i}_mod")
        else:
            batches[i]["B1V1"] = st.text_input("2h Viscosity", key=f"b{i}v1_far")
            batches[i]["B1V2"] = st.text_input("24h Viscosity", key=f"b{i}v2_far")
            batches[i]["PH"] = st.text_input("pH", key=f"ph{i}_far")
            batches[i]["MESH"] = st.text_input("200# Mesh", key=f"mesh{i}")
            batches[i]["BD"] = st.text_input("Bulk Density", key=f"bd{i}")
            batches[i]["F"] = st.text_input("Fann 3min", key=f"f{i}")
            batches[i]["FV"] = st.text_input("Fann 30min", key=f"fv{i}")

generate = st.button("Generate PSS + COA (merged PDF)")

if generate:
    if current_container > total_containers:
        st.error("Current container cannot be greater than total containers.")
    else:
        pss_path = find_template(PSS_TEMPLATES[doc_type])
        coa_path = find_template(COA_TEMPLATES[doc_type])

        missing = []
        if not pss_path:
            missing.append(PSS_TEMPLATES[doc_type])
        if not coa_path:
            missing.append(COA_TEMPLATES[doc_type])

        if missing:
            st.error(f"Missing template file(s): {', '.join(missing)}. "
                     "Make sure they are committed to the repo root.")
        else:
            try:
                date_str = date_picker.strftime("%d/%m/%Y")

                # ---------- PSS mapping ----------
                po_value = po_id.strip() if po_id and po_id.strip() else "PO012"
                pss_mapping = {
                    "{{DD/MM/YYYY}}": date_str,
                    "DD/MM/YYYY": date_str,
                    "{{PO012}}": po_value,
                    "{{B1}}": batches["1"]["BATCH"], "B1": batches["1"]["BATCH"],
                    "{{B2}}": batches["2"]["BATCH"], "B2": batches["2"]["BATCH"],
                    "{{B3}}": batches["3"]["BATCH"], "B3": batches["3"]["BATCH"],
                    "{{B4}}": batches["4"]["BATCH"], "B4": batches["4"]["BATCH"],
                }

                # ---------- COA mapping ----------
                coa_mapping = {
                    "DD/MM/YYYY": date_str,
                    "DD-MM-YYYY": date_str.replace("/", "-"),
                }
                for i in ("1", "2", "3", "4"):
                    coa_mapping[f"BATCH_{i}"] = batches[i].get("BATCH", "")
                    coa_mapping[f"M{i}"] = batches[i].get("M", "")
                    coa_mapping[f"B{i}V1"] = batches[i].get("B1V1", "")
                    coa_mapping[f"B{i}V2"] = batches[i].get("B1V2", "")
                    coa_mapping[f"PH{i}"] = batches[i].get("PH", "")
                    if doc_type == "FAR":
                        coa_mapping[f"MESH{i}"] = batches[i].get("MESH", "")
                        coa_mapping[f"BD{i}"] = batches[i].get("BD", "")
                        coa_mapping[f"F{i}"] = batches[i].get("F", "")
                        coa_mapping[f"FV{i}"] = batches[i].get("FV", "")

                with st.spinner("Filling templates..."):
                    pss_docx = fill_template(pss_path, pss_mapping, style_preserving=False)
                    coa_docx = fill_template(coa_path, coa_mapping, style_preserving=True)

                with st.spinner("Merging PSS + COA into one document..."):
                    merged_docx = merge_docx(pss_docx, coa_docx)

                code3 = extract_3digit_code(po_value)
                cur, tot = int(current_container), int(total_containers)
                base_name = f"PSS_COA_{doc_type}_LIPL_{code3}_{cur}_of_{tot}"

                st.session_state.merged_docx = merged_docx
                st.session_state.merged_docx_filename = f"{base_name}.docx"
                st.session_state.merged_pdf_base_name = base_name
                st.session_state.pss_docx = pss_docx
                st.session_state.coa_docx = coa_docx
                st.session_state.docx_names = (
                    f"PSS_{doc_type}_LIPL_{code3}_{cur}_of_{tot}.docx",
                    f"COA_{doc_type}_LIPL_{code3}_{cur}_of_{tot}.docx",
                )
                # a fresh generation invalidates any previously converted PDF
                st.session_state.merged_pdf = None
                st.success("Generated successfully.")
            except Exception as e:
                st.error(f"Failed to generate: {e}")

if st.session_state.get("merged_docx"):
    st.download_button(
        "📥 Download merged DOCX (PSS + COA)",
        st.session_state.merged_docx,
        file_name=st.session_state.merged_docx_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.markdown("---")
    st.caption("Optional: also need a PDF? Convert the merged document below "
               "(requires LibreOffice on the host).")
    if st.button("Convert merged document to PDF"):
        try:
            with st.spinner("Converting to PDF..."):
                merged_pdf = docx_bytes_to_pdf_bytes(
                    st.session_state.merged_docx, "merged_temp"
                )
            st.session_state.merged_pdf = merged_pdf
        except Exception as e:
            st.error(f"PDF conversion failed: {e}")

    if st.session_state.get("merged_pdf"):
        st.download_button(
            "📥 Download merged PDF (PSS + COA)",
            st.session_state.merged_pdf,
            file_name=f"{st.session_state.merged_pdf_base_name}.pdf",
            mime="application/pdf",
        )

    with st.expander("Download individual Word documents"):
        pss_name, coa_name = st.session_state.docx_names
        st.download_button(
            "Download PSS.docx",
            st.session_state.pss_docx,
            file_name=pss_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.download_button(
            "Download COA.docx",
            st.session_state.coa_docx,
            file_name=coa_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
